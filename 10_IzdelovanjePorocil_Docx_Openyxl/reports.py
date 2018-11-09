from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
import os.path

# list of all the paths we need
path = r"C:\Users\nucic\Desktop\PYTHON\X_IGRISCE_PYTHON\Reports"

# excel file, we read the data from
exFile = "Zvezek1.xlsx"

#Create new document
document = Document()

#Finds all the images in a certain directory!!
images = [f for f in os.listdir(path) if f.lower().endswith(".jpg") if os.path.isfile(os.path.join(path, f))]
# print (images)

#Create first page of the report!
header = document.add_heading("Poročilo izvedenih meritev", 1)
header.alignment = WD_ALIGN_PARAGRAPH.CENTER

header1 = document.add_paragraph('Poročilo o opravljenih meritvah ')
header1.alignment = 1

document.add_page_break()

wb = openpyxl.load_workbook(exFile)
ws = wb.get_sheet_by_name('List1')

for row in ws.iter_rows(row_offset=0):
    try:
        sifra = str((row[0]).value)
        mm = (row[1].value)
        vodotok = (row[2].value)
        pretok = (row[3].value)
        # print (sifra, mm, vodotok, pretok,slika)

        document.add_paragraph(sifra + " - "+vodotok+" - " + mm)
        document.add_paragraph("Izmerjen pretok: " + str(pretok) + "m3/s")

        for img in images:
            if img[:4] == sifra:
                print ("sifra: ", sifra)
                print ("img ", img)
                document.add_picture(os.path.join(path, img), width=Cm(13))

        document.add_page_break()

    except Exception as e:
        print(e)

document.save('demo.docx')